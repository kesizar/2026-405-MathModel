# -*- coding: utf-8 -*-
import fitz, os, sys

BASE = r'D:/HuaweiMoveData/Users/17669/Desktop/A 题'
OUT = r'D:/A题_work'

files = {
    'TJMML_A': '题目_TJMML_A.txt',
    '第一届天津市五校数学建模联赛格式规范': '格式规范.txt',
    '第一届天津市五校数学建模联赛人工智能工具使用规定': 'AI使用规定.txt',
}

for name, out in files.items():
    doc = fitz.open(os.path.join(BASE, name + '.pdf'))
    parts = []
    for i, p in enumerate(doc):
        parts.append('===== PAGE %d =====' % (i + 1))
        parts.append(p.get_text())
    path = os.path.join(OUT, out)
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))
    print('wrote', path, 'pages', doc.page_count)
    doc.close()

# Also dump docx
try:
    import docx
    d = docx.Document(os.path.join(BASE, 'A题补充解释.docx'))
    parts = []
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append(' | '.join(c.text for c in row.cells))
    with open(os.path.join(OUT, '补充解释.txt'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(parts))
    print('wrote 补充解释.txt')
except Exception as e:
    print('docx err', e)
print('ALL DONE')
